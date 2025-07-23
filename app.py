# app.py
import streamlit as st
import os
import tempfile
import fitz                     
import google.generativeai as genai
from datetime import datetime
import logging
import zipfile
import io
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ───────────────────────────── Configurações básicas ────────────────────────────
logging.basicConfig(level=logging.INFO)
st.set_page_config(page_title="Legal Case Analyzer", layout="wide")

# ------------------------------------------------------------------------------
# Funções utilitárias
# ------------------------------------------------------------------------------

def get_api_key():
    try:
        return st.secrets["GOOGLE_API_KEY"]
    except Exception:
        st.sidebar.error("Chave da API não encontrada em secrets.toml.")
        st.sidebar.info("Adicione GOOGLE_API_KEY no arquivo secrets.toml.")
        return None


def configure_gemini():
    api_key = get_api_key()
    if not api_key:
        return False
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        st.sidebar.error(f"Erro ao configurar Gemini: {e}")
        return False


def extract_text_from_pdf(pdf_file) -> str:
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    texto = []
    for i, page in enumerate(doc):
        page_text = page.get_text("text", sort=True).strip()
        if page_text:
            texto.append(f"\nPágina {i+1}\n{page_text}\n")
        else:
            texto.append(f"\nPágina {i+1}: (sem texto pesquisável)\n")
    return "".join(texto)


def call_gemini(prompt: str, context_text: str, model_name: str = "gemini-2.0-flash") -> str:
    """
    Envia o prompt + TODO o texto do processo ao modelo (sem limite hard-coded).
    Verifique quotas de token do modelo em uso.
    """
    try:
        model = genai.GenerativeModel(model_name)
        response = model.generate_content(f"{prompt}\n\nDocumentos do processo:\n{context_text}")
        return response.text if hasattr(response, "text") else str(response)
    except Exception as e:
        logging.exception("Erro na chamada à LLM")
        return f"Erro na chamada à LLM: {e}"


def create_docx_from_analysis(case_name: str, analyses: dict) -> io.BytesIO:
    """
    Cria um documento DOCX com todas as análises de um processo
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx não está disponível")
    
    try:
        doc = Document()
        
        # Título do documento
        title = doc.add_heading(f'Análise Jurídica - {case_name}', 0)
        doc.add_paragraph(f'Data de geração: {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
        doc.add_paragraph()
        
        # Adiciona cada análise
        for analysis_type, content in analyses.items():
            doc.add_heading(analysis_type, level=1)
            # Divide o conteúdo em parágrafos para melhor formatação
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            doc.add_page_break()
        
        # Salva em BytesIO
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return docx_buffer
    
    except Exception as e:
        raise Exception(f"Erro ao criar documento DOCX: {str(e)}")


def create_html_from_analysis(case_name: str, analyses: dict) -> str:
    """
    Alternativa: Cria um HTML bem formatado com todas as análises
    """
    html_content = f"""
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Análise Jurídica - {case_name}</title>
        <style>
            body {{
                font-family: 'Times New Roman', serif;
                line-height: 1.6;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                color: #333;
            }}
            h1 {{
                color: #2c3e50;
                border-bottom: 3px solid #3498db;
                padding-bottom: 10px;
            }}
            h2 {{
                color: #34495e;
                margin-top: 30px;
                border-left: 4px solid #3498db;
                padding-left: 15px;
            }}
            .header {{
                text-align: center;
                margin-bottom: 40px;
                border-bottom: 1px solid #bdc3c7;
                padding-bottom: 20px;
            }}
            .analysis-section {{
                margin-bottom: 40px;
                padding: 20px;
                background-color: #f8f9fa;
                border-radius: 8px;
                border-left: 4px solid #3498db;
            }}
            .date {{
                color: #7f8c8d;
                font-style: italic;
            }}
            p {{
                text-align: justify;
                margin-bottom: 15px;
            }}
            .page-break {{
                page-break-before: always;
            }}
            @media print {{
                body {{
                    max-width: none;
                    margin: 0;
                }}
                .analysis-section {{
                    break-inside: avoid;
                }}
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <h1>Análise Jurídica</h1>
            <h2>{case_name}</h2>
            <p class="date">Data de geração: {datetime.now().strftime("%d/%m/%Y às %H:%M")}</p>
        </div>
    """
    
    for i, (analysis_type, content) in enumerate(analyses.items()):
        page_break_class = "page-break" if i > 0 else ""
        html_content += f"""
        <div class="analysis-section {page_break_class}">
            <h2>{analysis_type}</h2>
            <div>
        """
        
        # Converte quebras de linha em parágrafos HTML
        paragraphs = content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # Substitui quebras de linha simples por <br>
                formatted_paragraph = paragraph.strip().replace('\n', '<br>')
                html_content += f"<p>{formatted_paragraph}</p>"
        
        html_content += """
            </div>
        </div>
        """
    
    html_content += """
    </body>
    </html>
    """
    
    return html_content


def create_zip_with_all_analyses(all_analyses: dict, format_type: str = "txt") -> io.BytesIO:
    """
    Cria um arquivo ZIP com todas as análises
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        for case_name, analyses in all_analyses.items():
            safe_case_name = "".join(c for c in case_name if c.isalnum() or c in (' ', '_', '-')).strip()
            
            if format_type == "txt":
                # Cria um arquivo TXT para cada processo
                txt_content = f"ANÁLISE JURÍDICA - {case_name.upper()}\n"
                txt_content += f"Data de geração: {datetime.now().strftime('%d/%m/%Y às %H:%M')}\n"
                txt_content += "=" * 80 + "\n\n"
                
                for analysis_type, content in analyses.items():
                    txt_content += f"{analysis_type.upper()}\n"
                    txt_content += "-" * len(analysis_type) + "\n"
                    txt_content += content + "\n\n"
                    txt_content += "=" * 80 + "\n\n"
                
                zip_file.writestr(f"{safe_case_name}_{timestamp}.txt", txt_content.encode('utf-8'))
                
            elif format_type == "docx":
                try:
                    # Cria um arquivo DOCX para cada processo
                    docx_buffer = create_docx_from_analysis(case_name, analyses)
                    zip_file.writestr(f"{safe_case_name}_{timestamp}.docx", docx_buffer.getvalue())
                except Exception as e:
                    # Se falhar, cria um arquivo de erro
                    error_content = f"Erro ao criar DOCX para {case_name}: {str(e)}\n\nConteúdo em texto:\n\n"
                    for analysis_type, content in analyses.items():
                        error_content += f"{analysis_type}:\n{content}\n\n"
                    zip_file.writestr(f"{safe_case_name}_{timestamp}_ERROR.txt", error_content.encode('utf-8'))
                    
            elif format_type == "html":
                # Cria um arquivo HTML para cada processo
                html_content = create_html_from_analysis(case_name, analyses)
                zip_file.writestr(f"{safe_case_name}_{timestamp}.html", html_content.encode('utf-8'))
    
    zip_buffer.seek(0)
    return zip_buffer

# ------------------------------------------------------------------------------
# Interface
# ------------------------------------------------------------------------------

st.title("Análise de Intimações – Múltiplos Processos")
st.markdown("""
Faça upload de até **10 PDFs** (cada PDF representa um processo diferente) e obtenha **três análises** para cada:

1. **Teor da intimação**  
2. **Análise do caso + próximos passos**  
3. **Petição sugerida**  

Você pode exportar todas as análises em lote como arquivos TXT ou DOCX.
""")

api_ready = configure_gemini()
if not api_ready:
    st.stop()

# Inicializa o estado da sessão
if 'all_analyses' not in st.session_state:
    st.session_state.all_analyses = {}

st.header("Upload de documentos")
files = st.file_uploader(
    "Arraste seus PDFs aqui (máximo 10 processos)",
    type="pdf",
    accept_multiple_files=True
)

# Limita a 10 arquivos
if files and len(files) > 10:
    st.error("⚠️ Máximo de 10 processos permitidos. Por favor, selecione até 10 arquivos.")
    files = files[:10]

# ------------------------------------------------------------------------------
# Processamento dos PDFs
# ------------------------------------------------------------------------------

if files:
    st.info(f"📄 {len(files)} processo(s) carregado(s)")
    
    # Limpa análises anteriores se novos arquivos foram carregados
    current_file_names = [f.name for f in files]
    if set(current_file_names) != set(st.session_state.all_analyses.keys()):
        st.session_state.all_analyses = {}
    
    # Botão para processar todos os casos
    if st.button("🔎 Analisar todos os processos", type="primary"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        total_steps = len(files) * 3  # 3 análises por arquivo
        current_step = 0
        
        prompts = {
            "Avaliação da Intimação": 
                "Avalie o teor da intimação. Identifique natureza, requisitos legais, prazos e implicações.",
            "Providências Recomendadas": 
                "Analise o processo e recomende ações específicas para responder adequadamente à intimação.",
            "Petição": 
                "Redija a petição completa (cabeçalho, qualificação, fatos, fundamentos jurídicos, pedidos, fechamento)."
        }
        
        for idx, pdf_file in enumerate(files, 1):
            case_name = pdf_file.name.replace('.pdf', '')
            status_text.text(f"Processando: {case_name}")
            
            # Extrai texto do PDF
            with st.spinner(f"Extraindo texto do processo {idx}/{len(files)}..."):
                temp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
                temp.write(pdf_file.read())
                temp.flush()
                
                pdf_file.seek(0)  # Reset file pointer
                extracted_text = extract_text_from_pdf(pdf_file)
                
                temp.close()
                os.unlink(temp.name)
            
            # Executa as 3 análises para este processo
            case_analyses = {}
            
            for analysis_name, prompt in prompts.items():
                current_step += 1
                progress_bar.progress(current_step / total_steps)
                status_text.text(f"Analisando {case_name}: {analysis_name}")
                
                with st.spinner(f"Executando: {analysis_name}..."):
                    resultado = call_gemini(prompt, extracted_text)
                    case_analyses[analysis_name] = resultado
            
            # Armazena as análises deste processo
            st.session_state.all_analyses[case_name] = case_analyses
        
        progress_bar.progress(1.0)
        status_text.text("✅ Todas as análises concluídas!")
        st.success(f"🎉 Análise de {len(files)} processo(s) finalizada com sucesso!")

    # ------------------------------------------------------------------------------
    # Exibição dos resultados
    # ------------------------------------------------------------------------------
    
    if st.session_state.all_analyses:
        st.header("📋 Resultados das Análises")
        
        # Cria abas para cada processo
        if len(st.session_state.all_analyses) > 1:
            case_tabs = st.tabs(list(st.session_state.all_analyses.keys()))
            
            for tab_idx, (case_name, analyses) in enumerate(st.session_state.all_analyses.items()):
                with case_tabs[tab_idx]:
                    st.subheader(f"📁 {case_name}")
                    
                    # Cria sub-abas para cada tipo de análise
                    analysis_tabs = st.tabs(list(analyses.keys()))
                    
                    for analysis_idx, (analysis_type, content) in enumerate(analyses.items()):
                        with analysis_tabs[analysis_idx]:
                            st.markdown(content)
                            
                            # Botão de download individual
                            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                            st.download_button(
                                f"💾 Download {analysis_type}",
                                data=content,
                                file_name=f"{case_name}_{analysis_type.lower().replace(' ', '_')}_{ts}.txt",
                                mime="text/plain",
                                key=f"download_{case_name}_{analysis_type}"
                            )
        else:
            # Se apenas um processo, exibe diretamente
            case_name, analyses = list(st.session_state.all_analyses.items())[0]
            st.subheader(f"📁 {case_name}")
            
            analysis_tabs = st.tabs(list(analyses.keys()))
            
            for analysis_idx, (analysis_type, content) in enumerate(analyses.items()):
                with analysis_tabs[analysis_idx]:
                    st.markdown(content)
                    
                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                    st.download_button(
                        f"💾 Download {analysis_type}",
                        data=content,
                        file_name=f"{case_name}_{analysis_type.lower().replace(' ', '_')}_{ts}.txt",
                        mime="text/plain",
                        key=f"download_{case_name}_{analysis_type}"
                    )
        
        # ------------------------------------------------------------------------------
        # Exportação em lote
        # ------------------------------------------------------------------------------
        
        st.header("📦 Exportação em Lote")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if st.button("📄 Exportar todos como TXT", type="secondary"):
                try:
                    zip_buffer = create_zip_with_all_analyses(st.session_state.all_analyses, "txt")
                    
                    st.download_button(
                        label="💾 Download ZIP com arquivos TXT",
                        data=zip_buffer.getvalue(),
                        file_name=f"analises_juridicas_txt_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip"
                    )
                except Exception as e:
                    st.error(f"Erro ao criar arquivos TXT: {str(e)}")
        
        with col2:
            docx_button_label = "📝 Exportar todos como DOCX"
            docx_button_disabled = False
            
            if not DOCX_AVAILABLE:
                docx_button_label += " (Indisponível)"
                docx_button_disabled = True
                st.warning("⚠️ python-docx não está instalado. Instale com: pip install python-docx")
            
            if st.button(docx_button_label, type="secondary", disabled=docx_button_disabled):
                try:
                    zip_buffer = create_zip_with_all_analyses(st.session_state.all_analyses, "docx")
                    
                    st.download_button(
                        label="💾 Download ZIP com arquivos DOCX",
                        data=zip_buffer.getvalue(),
                        file_name=f"analises_juridicas_docx_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip"
                    )
                except Exception as e:
                    st.error(f"Erro ao criar arquivos DOCX: {str(e)}")
        
        with col3:
            if st.button("🌐 Exportar todos como HTML", type="secondary"):
                try:
                    zip_buffer = create_zip_with_all_analyses(st.session_state.all_analyses, "html")
                    
                    st.download_button(
                        label="💾 Download ZIP com arquivos HTML",
                        data=zip_buffer.getvalue(),
                        file_name=f"analises_juridicas_html_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                        mime="application/zip"
                    )
                except Exception as e:
                    st.error(f"Erro ao criar arquivos HTML: {str(e)}")
        
        if DOCX_AVAILABLE:
            st.info("💡 Os arquivos serão compactados em um ZIP contendo uma análise completa para cada processo.")
        else:
            st.info("💡 Os arquivos serão compactados em um ZIP. Para DOCX, instale: pip install python-docx")

else:
    st.info("📤 Envie até 10 PDFs para começar (cada PDF representa um processo diferente).")
